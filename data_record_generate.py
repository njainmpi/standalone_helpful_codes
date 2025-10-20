# This cell creates a fully-local Flask portal that saves progress reports as .docx files
# and (optionally) integrates with a Kanban portal via a local webhook or JSON drop.
# It writes all necessary files and zips them for easy download.

import os, json, zipfile, textwrap, datetime, pathlib

base = pathlib.Path("/mnt/data/progress_portal")
templates = base / "templates"
static = base / "static"
reports = base / "reports"
kanban_outbox = base / "kanban_outbox"
base.mkdir(parents=True, exist_ok=True)
templates.mkdir(exist_ok=True)
static.mkdir(exist_ok=True)
reports.mkdir(exist_ok=True)
kanban_outbox.mkdir(exist_ok=True)

app_py = r'''#!/usr/bin/env python3
import os
from datetime import datetime
from zoneinfo import ZoneInfo
from pathlib import Path
import json

from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from docx import Document
from docx.shared import Pt, Inches

# ---------------- Configuration ----------------
APP_TITLE = "Naman's Local Progress Portal"
TIMEZONE = "Europe/Berlin"
SAVE_DIR = Path(__file__).parent / "reports"
SAVE_DIR.mkdir(exist_ok=True)
KANBAN_ENABLED_DEFAULT = False  # You can toggle this per-save via the form checkbox

# Filename style:
# - Default below matches your request "DD.MM.YY.docx"
# - If you actually meant 4-digit year, change to "DD.MM.YYYY"
FILENAME_STYLE = "DD.MM.YY"  # choices: "DD.MM.YY", "DD.MM.YYYY"

# Kanban integration:
# Option A: POST to a local Kanban API endpoint (e.g., your portal)
KANBAN_ENDPOINT = os.environ.get("KANBAN_ENDPOINT", "").strip()  # e.g., http://127.0.0.1:8000/api/cards
KANBAN_API_KEY = os.environ.get("KANBAN_API_KEY", "").strip()    # if your portal requires a token

# Option B: Drop JSON file to kanban_outbox directory for your portal to ingest
KANBAN_OUTBOX = Path(__file__).parent / "kanban_outbox"
KANBAN_OUTBOX.mkdir(exist_ok=True)

app = Flask(__name__)
app.secret_key = "local-only-secret"  # for flash messages (local app)

def format_filename(dt: datetime) -> str:
    if FILENAME_STYLE == "DD.MM.YY":
        return dt.strftime("%d.%m.%y") + ".docx"
    elif FILENAME_STYLE == "DD.MM.YYYY":
        return dt.strftime("%d.%m.%Y") + ".docx"
    else:
        # fallback
        return dt.strftime("%d.%m.%y") + ".docx"

def doc_title(dt: datetime) -> str:
    # Title inside the Word document
    return f"Naman's Progress Report for {dt.strftime('%d.%m.%y')}"

def add_heading(document: Document, text: str, level: int = 1):
    document.add_heading(text, level=level)

def add_bullets(document: Document, items):
    for it in items:
        if it and it.strip():
            document.add_paragraph(it.strip(), style='List Bullet')

def create_docx(payload: dict, tz: str = TIMEZONE) -> Path:
    # Determine report date
    try:
        report_date = datetime.fromisoformat(payload.get("report_date")).replace(tzinfo=None)
    except Exception:
        # fallback to now
        report_date = datetime.now(ZoneInfo(tz)).replace(tzinfo=None)
    # Build filename and title
    fname = format_filename(report_date)
    title_text = doc_title(report_date)

    document = Document()

    # Title
    title_para = document.add_paragraph()
    run = title_para.add_run(title_text)
    run.bold = True
    run.font.size = Pt(18)

    document.add_paragraph(f"Date: {report_date.strftime('%A, %d %B %Y')}")

    # Sections
    # Summary
    summary = payload.get("summary", "").strip()
    if summary:
        add_heading(document, "Summary", level=1)
        document.add_paragraph(summary)

    # Accomplishments (bulleted)
    accomplishments = payload.get("accomplishments", [])
    if any(a.strip() for a in accomplishments):
        add_heading(document, "Accomplishments", level=1)
        add_bullets(document, accomplishments)

    # Experiments
    experiments = payload.get("experiments", "").strip()
    if experiments:
        add_heading(document, "Experiments", level=1)
        document.add_paragraph(experiments)

    # Results / Observations
    results = payload.get("results", "").strip()
    if results:
        add_heading(document, "Results / Observations", level=1)
        document.add_paragraph(results)

    # Issues / Blockers
    blockers = payload.get("blockers", "").strip()
    if blockers:
        add_heading(document, "Issues / Blockers", level=1)
        document.add_paragraph(blockers)

    # Next Steps
    next_steps = payload.get("next_steps", "").strip()
    if next_steps:
        add_heading(document, "Next Steps", level=1)
        document.add_paragraph(next_steps)

    # Links
    links = payload.get("links", "").strip()
    if links:
        add_heading(document, "Links", level=1)
        for line in links.splitlines():
            if line.strip():
                document.add_paragraph(line.strip())

    # Hours
    hours = payload.get("hours", "").strip()
    if hours:
        add_heading(document, "Hours", level=1)
        document.add_paragraph(hours + " hours")

    # Save
    save_path = SAVE_DIR / fname
    document.save(save_path)
    return save_path

def push_to_kanban(payload: dict):
    # Prepare a concise card-like representation
    card = {
        "title": payload.get("kanban_title") or f"Progress: {payload.get('report_date','')}",
        "description": payload.get("summary", ""),
        "accomplishments": [a for a in payload.get("accomplishments", []) if a.strip()],
        "experiments": payload.get("experiments", ""),
        "results": payload.get("results", ""),
        "blockers": payload.get("blockers", ""),
        "next_steps": payload.get("next_steps", ""),
        "links": payload.get("links", ""),
        "report_date": payload.get("report_date", ""),
        "created_at": datetime.now(ZoneInfo(TIMEZONE)).isoformat(),
    }

    # Option A: POST to an endpoint if provided
    if KANBAN_ENDPOINT:
        try:
            import requests  # optional dependency
            headers = {"Content-Type": "application/json"}
            if KANBAN_API_KEY:
                headers["Authorization"] = f"Bearer {KANBAN_API_KEY}"
            r = requests.post(KANBAN_ENDPOINT, headers=headers, data=json.dumps(card), timeout=3)
            r.raise_for_status()
            return {"status": "ok", "method": "endpoint", "detail": r.text[:200]}
        except Exception as e:
            # Fall through to outbox
            pass

    # Option B: Write into outbox as JSON
    try:
        fname = f"kanban_{datetime.now(ZoneInfo(TIMEZONE)).strftime('%Y%m%d_%H%M%S')}.json"
        out_path = KANBAN_OUTBOX / fname
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(card, f, ensure_ascii=False, indent=2)
        return {"status": "ok", "method": "outbox", "detail": str(out_path)}
    except Exception as e:
        return {"status": "error", "detail": str(e)}

@app.route("/")
def index():
    return redirect(url_for("new_report"))

@app.route("/report/new", methods=["GET"])
def new_report():
    # default date = today in Berlin
    today = datetime.now(ZoneInfo(TIMEZONE)).date().isoformat()
    return render_template("form.html",
                           app_title=APP_TITLE,
                           default_date=today,
                           kanban_enabled_default=KANBAN_ENABLED_DEFAULT)

@app.route("/report", methods=["POST"])
def save_report():
    # Collect fields
    payload = {
        "report_date": request.form.get("report_date") or datetime.now(ZoneInfo(TIMEZONE)).date().isoformat(),
        "summary": request.form.get("summary",""),
        "accomplishments": request.form.getlist("accomplishments[]"),
        "experiments": request.form.get("experiments",""),
        "results": request.form.get("results",""),
        "blockers": request.form.get("blockers",""),
        "next_steps": request.form.get("next_steps",""),
        "links": request.form.get("links",""),
        "hours": request.form.get("hours",""),
        "kanban_title": request.form.get("kanban_title",""),
    }

    save_path = create_docx(payload)

    # Kanban integration (optional)
    if request.form.get("push_kanban") == "on":
        result = push_to_kanban(payload)
        if result.get("status") == "ok":
            flash(f"Kanban integration successful via {result.get('method')}.", "success")
        else:
            flash(f"Kanban integration failed: {result.get('detail')}", "error")

    # Offer download immediately as attachment
    return send_file(save_path, as_attachment=True, download_name=save_path.name)

if __name__ == "__main__":
    # Run locally
    app.run(host="127.0.0.1", port=5000, debug=False)
'''

form_html = r'''{% extends "base.html" %}
{% block content %}
<div class="container">
  <h1>{{ app_title }}</h1>
  <form method="post" action="{{ url_for('save_report') }}">
    <div class="grid">
      <label>Date
        <input type="date" name="report_date" value="{{ default_date }}" required>
      </label>
      <label>Hours worked
        <input type="number" name="hours" min="0" step="0.25" placeholder="e.g., 7.5">
      </label>
    </div>

    <label>Summary (one paragraph)</label>
    <textarea name="summary" rows="4" placeholder="High-level overview of what you accomplished today."></textarea>

    <label>Accomplishments (add bullet items)</label>
    <div id="accomplishments">
      <div class="row">
        <input type="text" name="accomplishments[]" placeholder="e.g., Fixed preprocessing script for motion correction">
        <button type="button" class="btn small" onclick="removeRow(this)">−</button>
      </div>
    </div>
    <button type="button" class="btn" onclick="addRow('accomplishments')">+ Add Item</button>

    <label>Experiments</label>
    <textarea name="experiments" rows="4" placeholder="Protocols run, parameters, sequences, etc."></textarea>

    <label>Results / Observations</label>
    <textarea name="results" rows="4" placeholder="Key figures, percent signal change, CBV estimates, anomalies..."></textarea>

    <label>Issues / Blockers</label>
    <textarea name="blockers" rows="3" placeholder="Anything blocking you (equipment, approvals, data quality)..."></textarea>

    <label>Next Steps</label>
    <textarea name="next_steps" rows="3" placeholder="What you'll do next."></textarea>

    <label>Links (one per line)</label>
    <textarea name="links" rows="3" placeholder="GitHub commits, datasets, figures, notes..."></textarea>

    <fieldset class="kanban">
      <legend>Kanban Integration (optional)</legend>
      <label class="inline">
        <input type="checkbox" name="push_kanban" {% if kanban_enabled_default %}checked{% endif %}>
        Push a summary to my Kanban
      </label>
      <label>Kanban Card Title
        <input type="text" name="kanban_title" placeholder="e.g., Daily progress for {{ default_date }}">
      </label>
      <p class="hint">
        This will either POST to <code>KANBAN_ENDPOINT</code> (if set) or drop a JSON file into <code>kanban_outbox/</code>.
      </p>
    </fieldset>

    <div class="actions">
      <button class="btn primary" type="submit">Save as Word (.docx)</button>
    </div>
  </form>

  {% with messages = get_flashed_messages(with_categories=true) %}
    {% if messages %}
      <div class="flashes">
        {% for category, message in messages %}
          <div class="flash {{ category }}">{{ message }}</div>
        {% endfor %}
      </div>
    {% endif %}
  {% endwith %}
</div>

<script>
function addRow(containerId) {
  const container = document.getElementById(containerId);
  const row = document.createElement('div');
  row.className = 'row';
  row.innerHTML = `<input type="text" name="accomplishments[]" placeholder="Add accomplishment">
                   <button type="button" class="btn small" onclick="removeRow(this)">−</button>`;
  container.appendChild(row);
}
function removeRow(btn) {
  const row = btn.parentNode;
  row.parentNode.removeChild(row);
}
</script>
{% endblock %}
'''

base_html = r'''<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{{ app_title }}</title>
  <link rel="stylesheet" href="{{ url_for('static', filename='style.css') }}">
</head>
<body>
  <main>
    {% block content %}{% endblock %}
  </main>
</body>
</html>
'''

style_css = r'''*{box-sizing:border-box}body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Ubuntu,Cantarell,Noto Sans,sans-serif;margin:0;background:#0f1320;color:#eaeef5}
main{max-width:920px;margin:40px auto;padding:24px;background:#151a2b;border:1px solid #1e2540;border-radius:16px;box-shadow:0 10px 30px rgba(0,0,0,.3)}
h1{margin-top:0;font-size:1.6rem}
.container label{display:block;margin:14px 0 6px;font-weight:600}
.container input[type=text], .container input[type=number], .container input[type=date], .container textarea{
  width:100%;padding:10px 12px;border:1px solid #2a335c;background:#0e1427;color:#eaeef5;border-radius:10px;outline:none}
.container textarea{resize:vertical}
.grid{display:grid;grid-template-columns:1fr 1fr;gap:16px}
.row{display:flex;gap:8px;margin:8px 0}
.btn{padding:8px 12px;border-radius:10px;border:1px solid #2a335c;background:#0e1427;color:#eaeef5;cursor:pointer}
.btn:hover{filter:brightness(1.2)}
.btn.small{padding:6px 10px}
.btn.primary{background:#2d5bff;border-color:#284fe0}
.actions{margin-top:16px}
fieldset.kanban{margin-top:24px;border:1px solid #2a335c;border-radius:12px;padding:12px}
legend{padding:0 6px}
label.inline{display:flex;align-items:center;gap:10px}
.hint{opacity:.8;font-size:.9rem}
.flashes{margin-top:16px}
.flash{padding:10px;border-radius:8px;margin:6px 0}
.flash.success{background:#16371f;border:1px solid #1f7a36}
.flash.error{background:#371616;border:1px solid #7a1f1f}
code{background:#0e1427;padding:2px 6px;border-radius:6px;border:1px solid #2a335c}
'''

requirements_txt = r'''Flask==3.0.3
python-docx==1.1.2
requests==2.32.3
tzdata==2024.1
'''

readme_md = r'''# Naman's Local Progress Portal

A fully local Flask app to fill in a structured daily progress report and save it as a Word (`.docx`) file.
It can also (optionally) push a summary to your **Kanban** portal either via a local HTTP endpoint or by dropping JSON files into `kanban_outbox/` for your portal to ingest.

## Quick Start

```bash
cd progress_portal
python3 -m venv .venv
source .venv/bin/activate  # Windows: .venv\Scripts\activate
pip install -r requirements.txt
python app.py
